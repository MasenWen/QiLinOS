#!/usr/bin/env python3
"""
创建测试表单文件

生成三种格式相同内容的表单文件：TXT、DOCX、PDF
"""

from pathlib import Path
import sys

# 添加项目根目录到Python路径
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))


def create_word_document():
    """创建Word格式的表单文档"""
    try:
        import docx
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 创建新文档
        doc = docx.Document()
        
        # 添加标题
        title = doc.add_heading('个人信息登记表', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加分隔线
        doc.add_paragraph('=' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 基本信息部分
        doc.add_heading('基本信息：', level=2)
        
        basic_fields = [
            '姓名：__________________',
            '性别：__________________', 
            '年龄：__________________',
            '出生日期：__________________',
            '身份证号：__________________'
        ]
        
        for field in basic_fields:
            doc.add_paragraph(field)
        
        # 联系方式部分
        doc.add_heading('联系方式：', level=2)
        
        contact_fields = [
            '手机号码：__________________',
            '电子邮箱：__________________',
            '联系地址：__________________',
            '邮政编码：__________________'
        ]
        
        for field in contact_fields:
            doc.add_paragraph(field)
            
        # 工作信息部分
        doc.add_heading('工作信息：', level=2)
        
        work_fields = [
            '工作单位：__________________',
            '职位/职务：__________________',
            '部门：__________________',
            '工作地址：__________________',
            '工作电话：__________________'
        ]
        
        for field in work_fields:
            doc.add_paragraph(field)
            
        # 教育背景部分
        doc.add_heading('教育背景：', level=2)
        
        education_fields = [
            '最高学历：__________________',
            '毕业院校：__________________',
            '所学专业：__________________',
            '毕业时间：__________________'
        ]
        
        for field in education_fields:
            doc.add_paragraph(field)
            
        # 其他信息部分
        doc.add_heading('其他信息：', level=2)
        
        other_fields = [
            '紧急联系人：__________________',
            '紧急联系电话：__________________',
            '特殊说明：__________________'
        ]
        
        for field in other_fields:
            doc.add_paragraph(field)
            
        # 签名部分
        doc.add_paragraph()
        doc.add_paragraph('申请人签名：__________________')
        doc.add_paragraph('申请日期：__________________')
        
        # 注意事项
        doc.add_paragraph()
        note = doc.add_paragraph('注：请用黑色钢笔或签字笔填写，字迹清晰。')
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 保存文档
        output_path = Path(__file__).parent / 'test_form.docx'
        doc.save(output_path)
        print(f"✅ Word文档已创建: {output_path}")
        return True
        
    except ImportError:
        print("❌ 需要安装python-docx: pip install python-docx")
        return False
    except Exception as e:
        print(f"❌ 创建Word文档失败: {e}")
        return False


def create_pdf_document():
    """创建PDF格式的表单文档"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import inch
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        # 尝试注册中文字体
        try:
            # Windows系统常见中文字体
            import platform
            if platform.system() == "Windows":
                pdfmetrics.registerFont(TTFont('SimSun', 'C:/Windows/Fonts/simsun.ttc'))
                font_name = 'SimSun'
            else:
                # 使用默认字体
                font_name = 'Helvetica'
        except:
            font_name = 'Helvetica'
        
        output_path = Path(__file__).parent / 'test_form.pdf'
        
        # 创建PDF文档
        c = canvas.Canvas(str(output_path), pagesize=A4)
        width, height = A4
        
        # 设置字体
        c.setFont(font_name, 16)
        
        # 添加标题
        title = "个人信息登记表"
        title_width = c.stringWidth(title, font_name, 16)
        c.drawString((width - title_width) / 2, height - 50, title)
        
        # 分隔线
        c.setFont(font_name, 12)
        separator = "=" * 50
        separator_width = c.stringWidth(separator, font_name, 12)
        c.drawString((width - separator_width) / 2, height - 80, separator)
        
        # 当前Y位置
        y_position = height - 120
        line_height = 20
        
        # 添加表单字段
        fields = [
            "基本信息：",
            "超维语义：__________________",
            "性别：__________________", 
            "年龄：__________________",
            "出生日期：__________________",
            "身份证号：__________________",
            "",
            "联系方式：",
            "手机号码：__________________",
            "电子邮箱：__________________",
            "联系地址：__________________",
            "邮政编码：__________________",
            "",
            "工作信息：",
            "工作单位：__________________",
            "职位/职务：__________________",
            "部门：__________________",
            "工作地址：__________________",
            "工作电话：__________________",
            "",
            "教育背景：",
            "最高学历：__________________",
            "毕业院校：__________________",
            "所学专业：__________________",
            "毕业时间：__________________",
            "",
            "其他信息：",
            "紧急联系人：__________________",
            "紧急联系电话：__________________",
            "特殊说明：__________________",
            "",
            "申请人签名：__________________",
            "申请日期：__________________",
            "",
            "注：请用黑色钢笔或签字笔填写，字迹清晰。"
        ]
        
        for field in fields:
            if field.endswith("：") and not field.startswith("姓名") and not field.startswith("手机"):
                # 这是一个章节标题
                c.setFont(font_name, 14)
                c.drawString(50, y_position, field)
                c.setFont(font_name, 12)
            elif field.startswith("注："):
                # 注意事项居中显示
                note_width = c.stringWidth(field, font_name, 12)
                c.drawString((width - note_width) / 2, y_position, field)
            elif field == "":
                # 空行
                pass
            else:
                # 普通字段
                c.drawString(70, y_position, field)
            
            y_position -= line_height
            
            # 如果页面空间不够，可以添加新页面
            if y_position < 50:
                c.showPage()
                y_position = height - 50
                c.setFont(font_name, 12)
        
        # 保存PDF
        c.save()
        print(f"✅ PDF文档已创建: {output_path}")
        return True
        
    except ImportError:
        print("❌ 需要安装reportlab: pip install reportlab")
        return False
    except Exception as e:
        print(f"❌ 创建PDF文档失败: {e}")
        return False


def main():
    """主函数"""
    print("🚀 开始创建测试表单文件...")
    
    # 确保data目录存在
    data_dir = Path(__file__).parent
    data_dir.mkdir(exist_ok=True)
    
    # TXT文件已经存在
    txt_path = data_dir / 'test_form.txt'
    if txt_path.exists():
        print(f"✅ TXT文档已存在: {txt_path}")
    
    # 创建Word文档
    word_success = create_word_document()
    
    # 创建PDF文档
    pdf_success = create_pdf_document()
    
    # 总结
    print("\n📋 创建结果总结:")
    print(f"📄 TXT文档: {'✅' if txt_path.exists() else '❌'}")
    print(f"📝 Word文档: {'✅' if word_success else '❌'}")
    print(f"📕 PDF文档: {'✅' if pdf_success else '❌'}")
    
    if word_success and pdf_success:
        print("\n🎉 所有测试文件创建完成！")
        print("📁 文件位置: out-data/")
        print("   - test_form.txt")
        print("   - test_form.docx") 
        print("   - test_form.pdf")
    else:
        print("\n⚠️ 部分文件创建失败，请检查依赖库安装:")
        if not word_success:
            print("   pip install python-docx")
        if not pdf_success:
            print("   pip install reportlab")


if __name__ == '__main__':
    main() 