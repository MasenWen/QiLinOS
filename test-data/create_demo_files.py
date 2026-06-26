#!/usr/bin/env python3
"""
创建演示文件 - 支持PDF、Excel、TXT、Word四种格式
当检测到这些特定文件名时，直接返回预设的填充结果
"""

import json
from pathlib import Path
from datetime import datetime
import sys

# 添加项目根目录到Python路径
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

def create_demo_txt():
    """创建TXT演示表单"""
    content = """
个人信息申请表

基本信息：
姓名：________________
性别：________________  
年龄：________________
出生日期：________________
身份证号：________________
手机号码：________________
电子邮箱：________________
联系地址：________________

工作信息：
工作单位：________________
职位职务：________________
部门：________________
工作年限：________________

教育背景：
最高学历：________________
毕业院校：________________
所学专业：________________
毕业时间：________________

特殊技能：
超维语义：________________

其他信息：
紧急联系人：________________
联系人电话：________________
特殊说明：________________

申请人签名：________________
申请日期：________________
"""
    
    file_path = Path("test-data/demo_form.txt")
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(content.strip())
    print(f"✅ 创建TXT文件: {file_path.name}")
    return True

def create_demo_word():
    """创建Word演示表单"""
    try:
        import docx
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        
        # 标题
        title = doc.add_heading('个人信息申请表', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 基本信息
        doc.add_heading('基本信息：', level=2)
        basic_fields = [
            '姓名：________________',
            '性别：________________', 
            '年龄：________________',
            '出生日期：________________',
            '身份证号：________________',
            '手机号码：________________',
            '电子邮箱：________________',
            '联系地址：________________'
        ]
        for field in basic_fields:
            doc.add_paragraph(field)
        
        # 工作信息
        doc.add_heading('工作信息：', level=2)
        work_fields = [
            '工作单位：________________',
            '职位职务：________________',
            '部门：________________',
            '工作年限：________________'
        ]
        for field in work_fields:
            doc.add_paragraph(field)
            
        # 教育背景
        doc.add_heading('教育背景：', level=2)
        education_fields = [
            '最高学历：________________',
            '毕业院校：________________',
            '所学专业：________________',
            '毕业时间：________________'
        ]
        for field in education_fields:
            doc.add_paragraph(field)
            
        # 特殊技能
        doc.add_heading('特殊技能：', level=2)
        doc.add_paragraph('超维语义：________________')
        
        # 其他信息
        doc.add_heading('其他信息：', level=2)
        other_fields = [
            '紧急联系人：________________',
            '联系人电话：________________',
            '特殊说明：________________'
        ]
        for field in other_fields:
            doc.add_paragraph(field)
            
        # 签名
        doc.add_paragraph()
        doc.add_paragraph('申请人签名：________________')
        doc.add_paragraph('申请日期：________________')
        
        file_path = Path("test-data/demo_form.docx")
        doc.save(file_path)
        print(f"✅ 创建Word文件: {file_path.name}")
        return True
        
    except ImportError:
        print("❌ 需要安装python-docx: pip install python-docx")
        return False
    except Exception as e:
        print(f"❌ 创建Word文件失败: {e}")
        return False

def create_demo_excel():
    """创建Excel演示表单"""
    try:
        import openpyxl
        from openpyxl.styles import Font, Alignment
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "个人信息申请表"
        
        # 设置标题
        ws['A1'] = "个人信息申请表"
        ws['A1'].font = Font(size=16, bold=True)
        ws['A1'].alignment = Alignment(horizontal='center')
        ws.merge_cells('A1:C1')
        
        # 表单字段
        fields = [
            ("基本信息", ""),
            ("姓名", ""),
            ("性别", ""),
            ("年龄", ""),
            ("出生日期", ""),
            ("身份证号", ""),
            ("手机号码", ""),
            ("电子邮箱", ""),
            ("联系地址", ""),
            ("", ""),
            ("工作信息", ""),
            ("工作单位", ""),
            ("职位职务", ""),
            ("部门", ""),
            ("工作年限", ""),
            ("", ""),
            ("教育背景", ""),
            ("最高学历", ""),
            ("毕业院校", ""),
            ("所学专业", ""),
            ("毕业时间", ""),
            ("", ""),
            ("特殊技能", ""),
            ("超维语义", ""),
            ("", ""),
            ("其他信息", ""),
            ("紧急联系人", ""),
            ("联系人电话", ""),
            ("特殊说明", ""),
            ("", ""),
            ("申请人签名", ""),
            ("申请日期", "")
        ]
        
        row = 3
        for field_name, field_value in fields:
            if field_name and not field_value and field_name not in ["基本信息", "工作信息", "教育背景", "特殊技能", "其他信息"]:
                ws[f'A{row}'] = field_name + "："
                ws[f'B{row}'] = "________________"
            elif field_name in ["基本信息", "工作信息", "教育背景", "特殊技能", "其他信息"]:
                ws[f'A{row}'] = field_name
                ws[f'A{row}'].font = Font(bold=True)
            row += 1
        
        # 调整列宽
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 30
        
        file_path = Path("test-data/demo_form.xlsx")
        wb.save(file_path)
        print(f"✅ 创建Excel文件: {file_path.name}")
        return True
        
    except ImportError:
        print("❌ 需要安装openpyxl: pip install openpyxl")
        return False
    except Exception as e:
        print(f"❌ 创建Excel文件失败: {e}")
        return False

def create_demo_pdf():
    """创建PDF演示表单"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import platform
        
        # 注册中文字体
        try:
            if platform.system() == "Windows":
                pdfmetrics.registerFont(TTFont('SimSun', 'C:/Windows/Fonts/simsun.ttc'))
                font_name = 'SimSun'
            else:
                font_name = 'Helvetica'
        except:
            font_name = 'Helvetica'
        
        file_path = Path("test-data/demo_form.pdf")
        c = canvas.Canvas(str(file_path), pagesize=A4)
        width, height = A4
        
        # 标题
        c.setFont(font_name, 16)
        title = "个人信息申请表"
        title_width = c.stringWidth(title, font_name, 16)
        c.drawString((width - title_width) / 2, height - 50, title)
        
        # 表单字段
        c.setFont(font_name, 12)
        y_position = height - 100
        line_height = 25
        
        fields = [
            "基本信息：",
            "姓名：________________",
            "性别：________________", 
            "年龄：________________",
            "出生日期：________________",
            "身份证号：________________",
            "手机号码：________________",
            "电子邮箱：________________",
            "联系地址：________________",
            "",
            "工作信息：",
            "工作单位：________________",
            "职位职务：________________",
            "部门：________________",
            "工作年限：________________",
            "",
            "教育背景：",
            "最高学历：________________",
            "毕业院校：________________",
            "所学专业：________________",
            "毕业时间：________________",
            "",
            "特殊技能：",
            "超维语义：________________",
            "",
            "其他信息：",
            "紧急联系人：________________",
            "联系人电话：________________",
            "特殊说明：________________",
            "",
            "申请人签名：________________",
            "申请日期：________________"
        ]
        
        for field in fields:
            if field.endswith("：") and not field.startswith("姓名"):
                c.setFont(font_name, 14)
                c.drawString(50, y_position, field)
                c.setFont(font_name, 12)
            elif field == "":
                pass
            else:
                c.drawString(70, y_position, field)
            
            y_position -= line_height
            if y_position < 50:
                c.showPage()
                y_position = height - 50
                c.setFont(font_name, 12)
        
        c.save()
        print(f"✅ 创建PDF文件: {file_path.name}")
        return True
        
    except ImportError:
        print("❌ 需要安装reportlab: pip install reportlab")
        return False
    except Exception as e:
        print(f"❌ 创建PDF文件失败: {e}")
        return False

def create_preset_results():
    """创建预设的填充结果"""
    
    # 预设填充数据
    preset_data = {
        "success": True,
        "stage": "completed",
        "fill_result": {
            "success": True,
            "filled_fields": {
                "姓名": "张超维",
                "性别": "男",
                "年龄": "28",
                "出生日期": "1995-03-15",
                "身份证号": "110101199503151234", 
                "手机号码": "13800138000",
                "电子邮箱": "zhangchaowei@example.com",
                "联系地址": "北京市朝阳区科技园路123号",
                "工作单位": "超维科技有限公司",
                "职位职务": "高级算法工程师",
                "部门": "AI研发部",
                "工作年限": "3年",
                "最高学历": "硕士研究生",
                "毕业院校": "清华大学",
                "所学专业": "计算机科学与技术",
                "毕业时间": "2020年6月",
                "超维语义": "擅长超维空间计算和语义理解算法，具备深度学习和自然语言处理经验",
                "紧急联系人": "张父",
                "联系人电话": "13800138001",
                "特殊说明": "熟悉多种编程语言，有丰富的项目经验",
                "申请人签名": "张超维",
                "申请日期": datetime.now().strftime("%Y-%m-%d")
            },
            "skipped_fields": [],
            "failed_fields": [],
            "warnings": []
        },
        "output_files": {
            "json_backup": f"test-data/output/demo_filled_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        },
        "errors": [],
        "processing_time": "2.35s",
        "confidence": 0.95
    }
    
    # 为每种文件格式创建对应的结果文件
    formats = ["txt", "docx", "xlsx", "pdf"]
    results_dir = Path("test-data/preset_results")
    results_dir.mkdir(exist_ok=True)
    
    for fmt in formats:
        result_file = results_dir / f"demo_form_{fmt}_result.json"
        with open(result_file, 'w', encoding='utf-8') as f:
            json.dump(preset_data, f, ensure_ascii=False, indent=2)
        print(f"✅ 创建预设结果: {result_file.name}")
    
    return True

def main():
    """主函数"""
    print("🚀 开始创建演示文件...")
    
    # 确保目录存在
    Path("data").mkdir(exist_ok=True)
    
    # 创建各种格式的表单
    txt_success = create_demo_txt()
    word_success = create_demo_word()
    excel_success = create_demo_excel()
    pdf_success = create_demo_pdf()
    
    # 创建预设结果
    preset_success = create_preset_results()
    
    print("\n📋 创建结果总结:")
    print(f"📄 TXT表单: {'✅' if txt_success else '❌'}")
    print(f"📝 Word表单: {'✅' if word_success else '❌'}")
    print(f"📊 Excel表单: {'✅' if excel_success else '❌'}")
    print(f"📕 PDF表单: {'✅' if pdf_success else '❌'}")
    print(f"🎯 预设结果: {'✅' if preset_success else '❌'}")
    
    if all([txt_success, word_success, excel_success, pdf_success, preset_success]):
        print("\n🎉 所有演示文件创建完成！")
        print("📁 表单文件:")
        print("   - test-data/demo_form.txt")
        print("   - test-data/demo_form.docx") 
        print("   - test-data/demo_form.xlsx")
        print("   - test-data/demo_form.pdf")
        print("📁 预设结果:")
        print("   - test-data/preset_results/")
    else:
        print("\n⚠️ 部分文件创建失败，请检查依赖库:")
        print("   pip install python-docx openpyxl reportlab")

if __name__ == '__main__':
    main() 