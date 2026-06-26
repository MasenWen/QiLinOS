"""
DOC/DOCX文档解析器

解析Microsoft Word文档内容，提取文本、表格、表单字段等信息
"""

from pathlib import Path
from typing import Dict, List, Any, Optional

try:
    import docx
    from docx import Document
    from docx.table import Table  
    from docx.text.paragraph import Paragraph
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import win32com.client
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False

from ..core import BaseParser, DocumentType, ProcessingStatus
from ..config import get_settings
from ..utils.logger import get_logger, performance_monitor, error_handler
from src.utils.db_manager import node_state

class DocParser(BaseParser):
    """DOC/DOCX文档解析器"""
    
    def __init__(self):
        super().__init__()
        self.logger = get_logger(__name__)
        self.settings = get_settings()
        
        if not HAS_DOCX:
            self.logger.warning("python-docx库未安装，某些功能可能不可用")
        
        if not HAS_WIN32COM:
            # self.logger.warning("win32com库未安装，无法处理.doc格式文件")
            pass
    
    def get_supported_formats(self) -> List[DocumentType]:
        """返回支持的文档格式列表"""
        return [DocumentType.DOC, DocumentType.DOCX]
    
    def validate_input(self, input_data: Any) -> bool:
        """验证输入数据的有效性"""
        if isinstance(input_data, (str, Path)):
            file_path = Path(input_data)
            return (file_path.exists() and 
                    file_path.suffix.lower() in ['.doc', '.docx'])
        return False
    
    def process(self, input_data: Any) -> Any:
        """处理输入数据并返回结果"""
        if isinstance(input_data, (str, Path)):
            return self.parse(Path(input_data))
        else:
            raise ValueError(f"不支持的输入类型: {type(input_data)}")
    
    @error_handler()
    @performance_monitor()
    def parse(self, file_path: Path) -> Dict[str, Any]:
        """
        解析DOC/DOCX文档
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            解析结果字典
        """
        self.logger.info(f"{node_state}-=-填表员===开始解析Word文档: {file_path}")
        
        try:
            result = {
                'file_path': str(file_path),
                'file_size': file_path.stat().st_size,
                'document_type': self._get_doc_type(file_path).value,
                'metadata': {},
                'paragraphs': [],
                'text_content': '',
                'tables': [],
                'forms': [],
                'styles': [],
                'images': [],
            }
            
            # 根据文件扩展名选择解析方法
            if file_path.suffix.lower() == '.docx':
                self._parse_docx(file_path, result)
            else:
                self._parse_doc(file_path, result)
            
            self.status = ProcessingStatus.SUCCESS
            self.logger.info(f"{node_state}-=-填表员===Word文档解析完成: {file_path}")
            
            return result
            
        except Exception as e:
            self.status = ProcessingStatus.FAILED
            self.logger.error(f"Word文档解析失败 {file_path}: {e}")
            raise
    
    def _parse_docx(self, file_path: Path, result: Dict[str, Any]):
        """解析DOCX格式文档"""
        if not HAS_DOCX:
            raise ImportError("python-docx库未安装，无法解析DOCX文件")
        
        try:
            doc = docx.Document(file_path)
            
            # 提取元数据
            result['metadata'] = self._extract_docx_metadata(doc)
            
            # 提取段落
            for para_idx, paragraph in enumerate(doc.paragraphs):
                para_info = {
                    'index': para_idx,
                    'text': paragraph.text,
                    'style': paragraph.style.name if paragraph.style else '',
                    'alignment': str(paragraph.alignment) if paragraph.alignment else '',
                    'runs': [],
                }
                
                # 提取段落中的运行（格式信息）
                for run in paragraph.runs:
                    run_info = {
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name if run.font.name else '',
                        'font_size': run.font.size.pt if run.font.size else 0,
                    }
                    para_info['runs'].append(run_info)
                
                result['paragraphs'].append(para_info)
                result['text_content'] += paragraph.text + '\n'
            
            # 提取表格
            for table_idx, table in enumerate(doc.tables):
                table_info = self._extract_docx_table(table, table_idx)
                result['tables'].append(table_info)
            
            # 提取样式
            result['styles'] = self._extract_docx_styles(doc)
            
            # 检测表单字段
            result['forms'] = self._detect_docx_form_fields(doc)

            print(result)
            
        except Exception as e:
            self.logger.error(f"解析DOCX文件失败: {e}")
            raise
    
    def _parse_doc(self, file_path: Path, result: Dict[str, Any]):
        """解析DOC格式文档（使用COM接口）"""
        if not HAS_WIN32COM:
            self.logger.warning("win32com库未安装，无法解析DOC文件，尝试转换后解析")
            # 这里可以尝试其他方法，比如使用在线转换服务
            return
        
        try:
            # 启动Word应用程序
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            
            try:
                # 打开文档
                doc = word_app.Documents.Open(str(file_path.absolute()))
                
                # 提取文本内容
                result['text_content'] = doc.Content.Text
                
                # 提取段落
                for para_idx, paragraph in enumerate(doc.Paragraphs):
                    para_info = {
                        'index': para_idx,
                        'text': paragraph.Range.Text.strip(),
                        'style': paragraph.Style.NameLocal if paragraph.Style else '',
                    }
                    result['paragraphs'].append(para_info)
                
                # 提取表格
                for table_idx, table in enumerate(doc.Tables):
                    table_info = self._extract_doc_table(table, table_idx)
                    result['tables'].append(table_info)
                
                # 关闭文档
                doc.Close()
                
            finally:
                # 退出Word应用程序
                word_app.Quit()
                
        except Exception as e:
            self.logger.error(f"使用COM接口解析DOC文件失败: {e}")
            # 尝试备用方法
            self._parse_doc_fallback(file_path, result)
    
    def _parse_doc_fallback(self, file_path: Path, result: Dict[str, Any]):
        """DOC文件解析的备用方法"""
        self.logger.info(f"{node_state}-=-填表员===尝试使用备用方法解析DOC文件")
        
        # 这里可以实现其他解析方法，比如：
        # 1. 使用python-docx2txt库
        # 2. 使用textract库
        # 3. 调用LibreOffice命令行工具
        
        try:
            # 尝试使用textract（如果安装了的话）
            import textract
            text = textract.process(str(file_path)).decode('utf-8')
            result['text_content'] = text
            result['paragraphs'] = [{'index': 0, 'text': text, 'style': ''}]
            
        except ImportError:
            self.logger.warning("textract库未安装，无法使用备用解析方法")
        except Exception as e:
            self.logger.warning(f"备用解析方法失败: {e}")
    
    def _extract_docx_metadata(self, doc: Document) -> Dict[str, Any]:
        """提取DOCX文档元数据"""
        metadata = {}
        
        try:
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'category': core_props.category or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'revision': str(core_props.revision) if core_props.revision else '',
            }
            
            # 清理空值
            metadata = {k: v for k, v in metadata.items() if v}
            
        except Exception as e:
            self.logger.warning(f"提取DOCX元数据失败: {e}")
        
        return metadata
    
    def _extract_docx_table(self, table: Table, table_idx: int) -> Dict[str, Any]:
        """提取DOCX表格信息"""
        table_info = {
            'index': table_idx,
            'rows': len(table.rows),
            'columns': len(table.columns) if table.rows else 0,
            'data': [],
            'style': table.style.name if table.style else '',
        }
        
        try:
            # 提取表格数据
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_info['data'].append(row_data)
            
        except Exception as e:
            self.logger.warning(f"提取表格数据失败: {e}")
        
        return table_info
    
    def _extract_doc_table(self, table: Any, table_idx: int) -> Dict[str, Any]:
        """提取DOC表格信息（COM接口）"""
        table_info = {
            'index': table_idx,
            'rows': table.Rows.Count,
            'columns': table.Columns.Count,
            'data': [],
        }
        
        try:
            # 提取表格数据
            for row_idx in range(1, table.Rows.Count + 1):
                row_data = []
                for col_idx in range(1, table.Columns.Count + 1):
                    try:
                        cell = table.Cell(row_idx, col_idx)
                        cell_text = cell.Range.Text.strip()
                        # 移除表格单元格结尾的特殊字符
                        cell_text = cell_text.replace('\r\x07', '')
                        row_data.append(cell_text)
                    except:
                        row_data.append('')
                table_info['data'].append(row_data)
                
        except Exception as e:
            self.logger.warning(f"提取DOC表格数据失败: {e}")
        
        return table_info
    
    def _extract_docx_styles(self, doc: Document) -> List[Dict[str, Any]]:
        """提取DOCX文档样式"""
        styles = []
        
        try:
            for style in doc.styles:
                style_info = {
                    'name': style.name,
                    'type': str(style.type),
                    'builtin': style.builtin,
                }
                styles.append(style_info)
                
        except Exception as e:
            self.logger.warning(f"提取文档样式失败: {e}")
        
        return styles
    
    def _detect_docx_form_fields(self, doc: Document) -> List[Dict[str, Any]]:
        """检测DOCX文档中的表单字段"""
        form_fields = []
        
        try:
            # 检测文本中的表单模式
            full_text = '\n'.join([para.text for para in doc.paragraphs])
            patterns = self._detect_form_patterns(full_text)
            
            # 转换为表单字段格式
            for pattern in patterns:
                form_field = {
                    'type': pattern['type'],
                    'label': pattern['matched_text'],
                    'position': pattern['position'],
                    'required': False,
                    'value': '',
                }
                form_fields.append(form_field)
                
        except Exception as e:
            self.logger.warning(f"检测表单字段失败: {e}")
        
        return form_fields
    
    def _detect_form_patterns(self, text: str) -> List[Dict[str, Any]]:
        """检测文本中的表单模式"""
        patterns = []
        
        import re
        
        # 常见表单字段模式
        field_patterns = {
            'name': [
                r'姓名[：:]\s*[_\s]*',
                r'Name[：:]\s*[_\s]*',
                r'名称[：:]\s*[_\s]*',
            ],
            'id_card': [
                r'身份证号[：:]\s*[_\s]*',
                r'ID[：:]\s*[_\s]*',
                r'证件号码[：:]\s*[_\s]*',
            ],
            'phone': [
                r'电话[：:]\s*[_\s]*',
                r'手机[：:]\s*[_\s]*',
                r'Phone[：:]\s*[_\s]*',
            ],
            'email': [
                r'邮箱[：:]\s*[_\s]*',
                r'Email[：:]\s*[_\s]*',
                r'电子邮件[：:]\s*[_\s]*',
            ],
            'address': [
                r'地址[：:]\s*[_\s]*',
                r'Address[：:]\s*[_\s]*',
                r'住址[：:]\s*[_\s]*',
            ],
            'date': [
                r'日期[：:]\s*[_\s]*',
                r'Date[：:]\s*[_\s]*',
                r'时间[：:]\s*[_\s]*',
            ],
        }
        
        for field_type, type_patterns in field_patterns.items():
            for pattern in type_patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    patterns.append({
                        'type': field_type,
                        'pattern': pattern,
                        'position': match.span(),
                        'matched_text': match.group(),
                    })
        
        return patterns
    
    def _get_doc_type(self, file_path: Path) -> DocumentType:
        """根据文件扩展名确定文档类型"""
        suffix = file_path.suffix.lower()
        if suffix == '.docx':
            return DocumentType.DOCX
        else:
            return DocumentType.DOC
    
    def get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名"""
        return ['.doc', '.docx']
    
    def get_status(self) -> ProcessingStatus:
        """获取处理状态"""
        return self.status 