"""
文档写入器

负责将填写的数据写入到原始文档格式中
"""

import shutil
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional, Union

from ..core import BaseProcessor, ProcessingStatus  
from ..config import get_settings
from ..utils.logger import get_logger, performance_monitor, error_handler
from autofill.llm.local_llm_client import LocalLLMClientAdapter
from src.utils.db_manager import node_state

class DocumentWriter(BaseProcessor):
    """文档写入器基类"""
    
    def __init__(self):
        super().__init__()
        self.logger = get_logger(__name__)
        self.settings = get_settings()
        self.status = ProcessingStatus.PENDING
    
    def process(self, input_data: Any) -> Any:
        """处理输入数据并返回结果"""
        # 这个方法在这里不适用，我们使用专门的write_filled_document方法
        return input_data
    
    def validate_input(self, input_data: Any) -> bool:
        """验证输入数据的有效性"""
        # 验证输入数据是否包含必要的字段
        if not isinstance(input_data, dict):
            return False
        
        required_keys = ['original_path', 'filled_data']
        return all(key in input_data for key in required_keys)
    
    @performance_monitor
    @error_handler()
    def write_filled_document(self, 
                                   original_path: Union[str, Path],
                                   filled_data: Dict[str, Any],
                                   output_path: Optional[Union[str, Path]] = None,
                                   form_analysis_result: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        将填写数据写入文档并保存
        
        Args:
            original_path: 原始文档路径
            filled_data: 填写的数据
            output_path: 输出路径（可选）
            form_analysis_result: 表单分析结果（包含字段信息）
            
        Returns:
            写入结果
        """
        original_path = Path(original_path)
        
        if not original_path.exists():
            raise FileNotFoundError(f"原始文档不存在: {original_path}")
        
        # 根据文件类型选择合适的写入器
        writer = self._get_writer_for_format(original_path)
        if not writer:
            raise ValueError(f"不支持的文档格式: {original_path.suffix}")
        
        # 生成输出路径
        if not output_path:
            output_path = self._generate_output_path(original_path)
        
        try:
            result = writer.write_document(original_path, filled_data, output_path, form_analysis_result)
            self.status = ProcessingStatus.SUCCESS
            return result
        except Exception as e:
            self.status = ProcessingStatus.FAILED
            self.logger.error(f"文档写入失败: {e}")
            raise
    
    def _get_writer_for_format(self, file_path: Path) -> Optional['BaseDocumentWriter']:
        """根据文件格式获取对应的写入器"""
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return PDFWriter()
        elif suffix in ['.docx', '.doc']:
            return WordWriter()
        elif suffix in ['.xlsx', '.xls']:
            return ExcelWriter()
        elif suffix == '.txt':
            return TextWriter()
        else:
            # 对于不支持的格式，使用通用文本写入器作为后备
            return FallbackWriter()
    
    def _generate_output_path(self, original_path: Path) -> Path:
        """生成输出文件路径"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = Path(self.settings.system.data_dir) / "output"
        output_dir.mkdir(exist_ok=True)
        
        stem = original_path.stem
        suffix = original_path.suffix
        output_name = f"{stem}_filled_{timestamp}{suffix}"
        
        return output_dir / output_name


class BaseDocumentWriter:
    """文档写入器基类"""
    
    def __init__(self):
        self.logger = get_logger(self.__class__.__name__)
    
    def write_document(self, 
                           original_path: Path,
                           filled_data: Dict[str, Any],
                           output_path: Path,
                           form_analysis_result: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """写入文档"""
        raise NotImplementedError
    
    def _get_field_name_variants(self, field_name, field_label=None):
        """
        获取字段名的智能变体映射
        处理LightRAG返回字段名与实际表单字段名的常见差异
        """
        variants = []
        
        if not field_name:
            return variants
        
        # 定义常见的字段名映射规则
        field_mapping = {
            # 手机相关
            "手机号码": ["手机号", "电话", "电话号码", "手机", "联系电话", "联系方式", "电话号"],
            "手机号": ["手机号码", "电话", "电话号码", "联系电话"],
            "电话": ["手机号码", "手机号", "电话号码", "联系电话", "联系方式"],
            "联系电话": ["手机号码", "手机号", "电话", "电话号码"],
            
            # 邮箱相关
            "电子邮箱": ["邮箱", "邮件", "电邮", "邮件地址", "电子邮件", "E-mail", "Email"],
            "邮箱": ["电子邮箱", "邮件", "电邮", "邮件地址", "电子邮件"],
            "邮件": ["电子邮箱", "邮箱", "邮件地址", "电子邮件"],
            "电子邮件": ["电子邮箱", "邮箱", "邮件", "邮件地址"],
            
            # 地址相关
            "联系地址": ["地址", "家庭住址", "住址", "通讯地址", "现住址", "居住地址"],
            "地址": ["联系地址", "家庭住址", "住址", "通讯地址", "现住址"],
            "住址": ["联系地址", "地址", "家庭住址", "通讯地址", "现住址"],
            "家庭住址": ["联系地址", "地址", "住址", "通讯地址"],
            
            # 姓名相关
            "姓名": ["真实姓名", "全名", "用户名", "name"],
            "真实姓名": ["姓名", "全名"],
            
            # 工作相关
            "工作单位": ["单位", "公司", "企业", "工作机构", "所在单位", "单位名称"],
            "单位": ["工作单位", "公司", "企业", "工作机构"],
            "公司": ["工作单位", "单位", "企业", "工作机构"],
            "工作年限": ["工作经验", "从业年限", "工龄", "工作时间", "从业经验"],
            "工作经验": ["工作年限", "从业年限", "工龄"],
            "职位职务": ["职位", "职务", "岗位", "职称"],
            "职位": ["职位职务", "职务", "岗位"],
            "职务": ["职位职务", "职位", "岗位"],
            
            # 身份证相关
            "身份证号": ["身份证", "身份证号码", "证件号", "证件号码", "ID号"],
            "身份证号码": ["身份证号", "身份证", "证件号", "证件号码"],
            
            # 学历相关
            "最高学历": ["学历", "教育程度", "文化程度"],
            "学历": ["最高学历", "教育程度", "文化程度"],
            "毕业院校": ["学校", "毕业学校", "院校", "毕业院校名称"],
            "所学专业": ["专业", "专业名称", "学习专业"],
            "专业": ["所学专业", "专业名称", "学习专业"],
            
            # 其他常见字段
            "性别": ["gender"],
            "年龄": ["age"],
            "出生日期": ["出生年月", "生日", "出生时间"],
            "紧急联系人": ["联系人", "紧急联系人姓名", "应急联系人"],
            "联系人电话": ["联系人手机", "联系人号码", "紧急联系电话"],
        }
        
        # 获取当前字段的所有变体
        if field_name in field_mapping:
            variants.extend(field_mapping[field_name])
        
        # 反向查找：如果当前字段名是某个映射的变体，也添加原字段名
        for main_field, variant_list in field_mapping.items():
            if field_name in variant_list and main_field not in variants:
                variants.append(main_field)
                # 同时添加该主字段的其他变体
                variants.extend([v for v in variant_list if v != field_name])
        
        # 处理标签提供的额外信息
        if field_label and field_label != field_name:
            if field_label in field_mapping:
                variants.extend(field_mapping[field_label])
        
        # 去重并过滤空值
        variants = list(dict.fromkeys([v for v in variants if v and v.strip()]))
        
        self.logger.debug(f"字段 '{field_name}' 的变体映射: {variants}")
        return variants


class PDFWriter(BaseDocumentWriter):
    """PDF文档写入器"""
    
    def write_document(self, 
                           original_path: Path,
                           filled_data: Dict[str, Any],
                           output_path: Path,
                           form_analysis_result: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """写入PDF文档"""
        try:
            # 检查依赖
            missing_deps = []
            try:
                import PyPDF2
            except ImportError:
                missing_deps.append("PyPDF2")
            
            try:
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter
            except ImportError:
                missing_deps.append("reportlab")
            
            if missing_deps:
                error_msg = f"PDF写入需要安装依赖: {', '.join(missing_deps)}"
                self.logger.error(error_msg)
                return {
                    'success': False,
                    'error': error_msg,
                    'method': 'pdf_form_fill',
                    'missing_dependencies': missing_deps
                }
            
            filled_fields = filled_data.get('filled_fields', {})
            
            if not filled_fields:
                return {
                    'success': False,
                    'error': '没有需要填写的数据',
                    'method': 'pdf_form_fill'
                }
            
            # 对于PDF，我们有几种策略：
            # 1. 如果是表单PDF，直接填写表单字段
            # 2. 如果是普通PDF，创建覆盖层
            
            result = self._fill_pdf_form_fields(original_path, filled_fields, output_path)
            
            if not result['success']:
                # 如果表单字段填写失败，尝试覆盖层方法
                self.logger.info(f"{node_state}-=-填表员===PDF表单字段填写失败，尝试覆盖层方法: {result.get('reason', 'Unknown error')}")
                result = self._create_pdf_overlay(original_path, filled_fields, output_path)
                
                if not result['success']:
                    # 如果覆盖层也失败，使用简单复制+文本附加方法
                    self.logger.info(f"{node_state}-=-填表员===PDF覆盖层创建失败，使用简单复制方法: {result.get('error', 'Unknown error')}")
                    result = self._create_simple_pdf_copy(original_path, filled_fields, output_path)
            
            return result
            
        except Exception as e:
            self.logger.error(f"PDF写入失败: {e}")
            return {
                'success': False,
                'output_file': None,
                'error': str(e),
                'method': 'pdf_form_fill'
            }
    
    def _fill_pdf_form_fields(self, 
                                   original_path: Path,
                                   filled_fields: Dict[str, Any],
                                   output_path: Path) -> Dict[str, Any]:
        """填写PDF表单字段"""
        try:
            import PyPDF2
            
            # 读取原始PDF
            with open(original_path, 'rb') as input_file:
                pdf_reader = PyPDF2.PdfReader(input_file)
                pdf_writer = PyPDF2.PdfWriter()
                
                # 检查是否有表单字段
                if not hasattr(pdf_reader, 'get_form_text_fields') and not hasattr(pdf_reader.pages[0], 'get'):
                    return {
                        'success': False,
                        'reason': 'PDF没有可填写的表单字段'
                    }
                
                # 填写字段
                for page in pdf_reader.pages:
                    pdf_writer.add_page(page)
                
                # 更新表单字段值
                if hasattr(pdf_writer, 'update_page_form_field_values'):
                    for field_name, field_value in filled_fields.items():
                        try:
                            pdf_writer.update_page_form_field_values(
                                pdf_writer.pages[0], {field_name: str(field_value)}
                            )
                        except Exception as e:
                            self.logger.warning(f"无法填写字段 {field_name}: {e}")
                
                # 保存填写后的PDF
                output_path.parent.mkdir(parents=True, exist_ok=True)
                with open(output_path, 'wb') as output_file:
                    pdf_writer.write(output_file)
                
                return {
                    'success': True,
                    'output_file': str(output_path),
                    'method': 'form_field_fill',
                    'filled_count': len(filled_fields)
                }
                
        except Exception as e:
            self.logger.warning(f"PDF表单字段填写失败: {e}")
            return {
                'success': False,
                'reason': str(e)
            }
    
    def _create_pdf_overlay(self, 
                                 original_path: Path,
                                 filled_fields: Dict[str, Any],
                                 output_path: Path) -> Dict[str, Any]:
        """创建PDF覆盖层"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            import PyPDF2
            
            # 创建临时覆盖层PDF
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_path = Path(temp_file.name)
                
                # 创建覆盖层
                c = canvas.Canvas(str(temp_path), pagesize=letter)
                
                # 添加文本（这里需要根据字段位置信息来定位）
                y_position = 750  # 从页面顶部开始
                for field_name, field_value in filled_fields.items():
                    c.drawString(100, y_position, f"{field_name}: {field_value}")
                    y_position -= 30
                
                c.save()
                
                # 合并原始PDF和覆盖层
                with open(original_path, 'rb') as orig_file, open(temp_path, 'rb') as overlay_file:
                    orig_pdf = PyPDF2.PdfReader(orig_file)
                    overlay_pdf = PyPDF2.PdfReader(overlay_file)
                    
                    pdf_writer = PyPDF2.PdfWriter()
                    
                    # 合并第一页
                    if len(orig_pdf.pages) > 0 and len(overlay_pdf.pages) > 0:
                        page = orig_pdf.pages[0]
                        page.merge_page(overlay_pdf.pages[0])
                        pdf_writer.add_page(page)
                        
                        # 添加其余页面
                        for i in range(1, len(orig_pdf.pages)):
                            pdf_writer.add_page(orig_pdf.pages[i])
                    
                    # 保存结果
                    output_path.parent.mkdir(parents=True, exist_ok=True)
                    with open(output_path, 'wb') as output_file:
                        pdf_writer.write(output_file)
                
                # 清理临时文件
                temp_path.unlink()
                
                return {
                    'success': True,
                    'output_file': str(output_path),
                    'method': 'overlay_merge',
                    'filled_count': len(filled_fields)
                }
                
        except Exception as e:
            self.logger.error(f"PDF覆盖层创建失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'method': 'overlay_merge'
            }

    def _create_simple_pdf_copy(self, 
                                     original_path: Path,
                                     filled_fields: Dict[str, Any],
                                     output_path: Path) -> Dict[str, Any]:
        """简单复制PDF并添加文本信息"""
        try:
            import shutil
            
            # 简单复制原始PDF
            output_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(original_path, output_path)
            
            # 创建一个伴随的文本文件包含填写信息
            text_output_path = output_path.with_suffix('.pdf_filled.txt')
            
            content = f"# PDF表单填写信息\n"
            content += f"# 原始文件: {original_path.name}\n"
            content += f"# 填写时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
            
            for field_name, field_value in filled_fields.items():
                content += f"{field_name}: {field_value}\n"
            
            with open(text_output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return {
                'success': True,
                'output_file': str(output_path),
                'method': 'simple_copy_with_text',
                'filled_count': len(filled_fields),
                'text_file': str(text_output_path)
            }
            
        except Exception as e:
            self.logger.error(f"PDF简单复制失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'method': 'simple_copy_with_text'
            }


class WordWriter(BaseDocumentWriter):
    """Word文档写入器"""
    
    def write_document(self, 
                           original_path: Path,
                           filled_data: Dict[str, Any],
                           output_path: Path,
                           form_analysis_result: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """写入Word文档"""
        try:
            # 检查依赖
            try:
                import docx
                from docx import Document
            except ImportError:
                # 如果没有python-docx，使用简单复制方法
                return self._create_simple_word_copy(original_path, filled_data, output_path)
            
            filled_fields = filled_data.get('filled_fields', {})
            
            if not filled_fields:
                return {
                    'success': False,
                    'error': '没有需要填写的数据',
                    'method': 'word_document'
                }
            
            try:
                # 复制原始文档
                shutil.copy2(original_path, output_path)
                
                # 打开文档
                doc = Document(output_path)
                
                filled_count = 0
                
                # 方法1: 替换文档中的占位符
                filled_count += self._replace_word_placeholders(doc, filled_fields, form_analysis_result)
                
                # 方法2: 填写表单控件
                filled_count += self._fill_word_form_controls(doc, filled_fields)
                
                # 方法3: 在指定位置插入内容
                filled_count += self._insert_word_content(doc, filled_fields)
                
                # 保存文档
                doc.save(output_path)
                
                return {
                    'success': True,
                    'output_file': str(output_path),
                    'method': 'word_document',
                    'filled_count': filled_count
                }
                
            except Exception as e:
                self.logger.warning(f"Word文档处理失败，使用简单复制方法: {e}")
                return self._create_simple_word_copy(original_path, filled_data, output_path)
            
        except Exception as e:
            self.logger.error(f"Word文档写入失败: {e}")
            return {
                'success': False,
                'output_file': None,
                'error': str(e),
                'method': 'word_document'
            }

    def _create_simple_word_copy(self, 
                                      original_path: Path,
                                      filled_data: Dict[str, Any],
                                      output_path: Path) -> Dict[str, Any]:
        """简单复制Word文档并添加文本信息"""
        try:
            import shutil
            
            filled_fields = filled_data.get('filled_fields', {})
            
            if not filled_fields:
                return {
                    'success': False,
                    'error': '没有需要填写的数据',
                    'method': 'simple_word_copy'
                }
            
            # 简单复制原始文档
            output_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(original_path, output_path)
            
            # 创建一个伴随的文本文件包含填写信息
            text_output_path = output_path.with_suffix('.docx_filled.txt')
            
            content = f"# Word文档填写信息\n"
            content += f"# 原始文件: {original_path.name}\n"
            content += f"# 填写时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
            
            for field_name, field_value in filled_fields.items():
                content += f"{field_name}: {field_value}\n"
            
            with open(text_output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return {
                'success': True,
                'output_file': str(output_path),
                'method': 'simple_word_copy',
                'filled_count': len(filled_fields),
                'text_file': str(text_output_path)
            }
            
        except Exception as e:
            self.logger.error(f"Word简单复制失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'method': 'simple_word_copy'
            }
    
    def _replace_word_placeholders(self, doc, filled_fields: Dict[str, Any], form_analysis_result: Optional[Dict[str, Any]] = None) -> int:
        """替换Word文档中的占位符"""
        filled_count = 0
        
        # 处理LightRAG字段数据格式
        processed_fields = {}
        for field_name, field_data in filled_fields.items():
            if isinstance(field_data, dict) and 'value' in field_data:
                # LightRAG格式：{'value': '值', 'method': '...', ...}
                processed_fields[field_name] = field_data['value']
            else:
                # 简单格式：直接是值
                processed_fields[field_name] = str(field_data)
        
        # 从表单分析结果中提取字段信息
        form_analysis_fields = []
        if form_analysis_result and isinstance(form_analysis_result, dict):
            form_analysis_fields = form_analysis_result.get('fields', [])
        
        # 一次性处理整个文档内容，避免重复LLM调用
        if filled_count == 0:  # 只在第一次时进行LLM全局填充
            # 提取整个文档的文本内容
            full_document_text = '\n'.join([p.text for p in doc.paragraphs])
            
            # 使用LLM一次性处理整个文档
            filled_document_text, global_fill_count = self._smart_form_fill(full_document_text, processed_fields, form_analysis_fields)
            
            if global_fill_count > 0:
                # 将LLM处理后的文本重新分配回段落
                filled_lines = filled_document_text.split('\n')
                for i, paragraph in enumerate(doc.paragraphs):
                    if i < len(filled_lines):
                        paragraph.text = filled_lines[i]
                filled_count += global_fill_count
                self.logger.info(f"{node_state}-=-填表员===✅ LLM全局文档填充完成，处理了整个文档，跳过后续填充步骤")
                
                # LLM全局填充成功，直接返回，避免重复处理
                return filled_count
            else:
                # 如果LLM填充失败，fallback到逐段落处理
                self.logger.warning("LLM全局填充失败，回退到段落级处理")
                for paragraph in doc.paragraphs:
                    original_text = paragraph.text
                    modified_text = original_text
                    
                    # 使用简单的直接字段匹配而不是LLM
                    modified_text, count = self._direct_field_mapping_fill(modified_text, processed_fields, form_analysis_fields)
                    filled_count += count
            
                    
                    # 如果文本有变化，更新段落
                    if modified_text != original_text:
                        paragraph.text = modified_text
        
        # 然后对整个文档尝试标准占位符替换（作为补充）
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            modified_text = original_text
            
            for field_name, field_value in processed_fields.items():
                # 查找常见的占位符格式
                placeholders = [
                    f"{{{field_name}}}",
                    f"[{field_name}]",
                    f"__{field_name}__",
                    f"{field_name}:",
                    f"{field_name}：",
                    # 增加更多常见格式
                    f"({field_name})",
                    f"<{field_name}>",
                    f"${{{field_name}}}",
                    f"%{field_name}%",
                ]
                
                for placeholder in placeholders:
                    if placeholder in modified_text:
                        modified_text = modified_text.replace(placeholder, str(field_value))
                        filled_count += 1
                        self.logger.debug(f"Word标准占位符替换: {placeholder} -> {field_value}")
                        break
            
            # 如果文本有变化，更新段落
            if modified_text != original_text:
                paragraph.text = modified_text
        
        # 在表格中查找和替换（使用直接字段匹配，避免重复LLM调用）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    modified_text = original_text
                    
                    # 使用直接字段匹配而不是LLM（避免重复调用）
                    modified_text, count = self._direct_field_mapping_fill(modified_text, processed_fields, form_analysis_fields)
                    filled_count += count
                    
                    # 然后尝试标准占位符
                    for field_name, field_value in processed_fields.items():
                        placeholders = [
                            f"{{{field_name}}}",
                            f"[{field_name}]",
                            f"__{field_name}__",
                            f"({field_name})",
                            f"<{field_name}>",
                        ]
                        
                        for placeholder in placeholders:
                            if placeholder in modified_text:
                                modified_text = modified_text.replace(placeholder, str(field_value))
                                filled_count += 1
                                self.logger.debug(f"Word表格标准占位符替换: {placeholder} -> {field_value}")
                                break
                    
                    # 如果文本有变化，更新单元格
                    if modified_text != original_text:
                        cell.text = modified_text
        
        return filled_count
    
    def _smart_form_fill(self, text: str, processed_fields: Dict[str, Any], form_analysis_fields: List[Dict[str, Any]] = None) -> tuple:
        """智能表单填写：使用直接字段匹配填充，跳过LLM处理以避免JSON解析错误"""
        
        # 如果没有字段数据，直接返回
        if not processed_fields:
            return text, 0
        
        # 直接使用字段匹配填充，不使用LLM（避免JSON解析失败）
        self.logger.info(f"{node_state}-=-填表员===🔄 已禁用LLM智能填充，使用直接关键字匹配模式")
        return self._direct_field_mapping_fill(text, processed_fields, form_analysis_fields or [])
    
    def _llm_global_document_fill(self, text: str, processed_fields: Dict[str, Any], form_analysis_fields: List[Dict[str, Any]]) -> tuple:
        """LLM全局文档填充：让LLM分析文档并返回结构化的填充指令，然后本地处理"""
        
        try:
            self.logger.info(f"{node_state}-=-填表员===开始LLM全局文档填充，处理 {len(processed_fields)} 个字段")
            
            # 构建LLM分析的prompt - 让LLM返回填充指令而不是完整文档
            analysis_prompt = self._build_llm_analysis_prompt(text, processed_fields, form_analysis_fields)
            
            # 调用LLM进行分析

            llm_client = LocalLLMClientAdapter()
            
            self.logger.debug("调用LLM进行文档分析和填充指令生成...")
            llm_response = llm_client.simple_chat(analysis_prompt)
            
            # 调试：记录LLM的原始响应
            self.logger.debug(f"🔍 LLM原始响应长度: {len(llm_response) if llm_response else 0}")
            if llm_response:
                self.logger.debug(f"🔍 LLM响应前200字符: {llm_response[:200]}")
            
            if not llm_response or llm_response.strip() == "":
                self.logger.error("LLM返回的分析结果为空")
                return text, 0
            
            # 解析LLM返回的填充指令 - 带重试机制
            fill_instructions = self._parse_llm_fill_instructions_with_retry(
                llm_response, text, processed_fields, form_analysis_fields, max_retries=2
            )
            
            if not fill_instructions:
                self.logger.warning("⚠️ LLM多次尝试后仍未返回有效指令，fallback到直接字段匹配模式")
                # 如果多次重试后仍失败，fallback到直接字段匹配
                return self._direct_field_mapping_fill(text, processed_fields, form_analysis_fields)
            
            # 根据填充指令执行本地填充
            filled_content, filled_count = self._execute_fill_instructions(text, fill_instructions)
            
            self.logger.info(f"{node_state}-=-填表员===LLM全局文档填充完成，成功填充 {filled_count} 个字段")
            
            return filled_content, filled_count
            
        except Exception as e:
            self.logger.error(f"LLM全局文档填充失败: {str(e)}")
            # 失败时返回原文本
            return text, 0
    
    def _build_llm_analysis_prompt(self, original_text: str, processed_fields: Dict[str, Any], form_analysis_fields: List[Dict[str, Any]]) -> str:
        """构建LLM分析的prompt，让LLM返回结构化的填充指令"""
        
        # 提取要填充的数据
        fill_data = {}
        for field_name, field_data in processed_fields.items():
            if isinstance(field_data, dict) and 'value' in field_data:
                fill_data[field_name] = field_data['value']
            else:
                fill_data[field_name] = field_data
        
        prompt = f"""⚠️ 重要：你必须严格按照JSON格式返回填充指令，不要直接修改文档内容！

【任务目标】
分析文档中的空白字段，生成精确的替换指令供程序执行。

【可用数据】
{self._format_fill_data_for_llm(fill_data)}

【原始文档内容】
{original_text}

【分析步骤】
1. 找到文档中形如"字段名：________________"的空白字段
2. 根据字段名的语义，从可用数据中找到对应的值
3. 生成replacement指令：保留字段名和冒号，只替换下划线部分

【严格的JSON输出格式】
你必须且只能返回以下格式的JSON，不要添加任何解释文字：

```json
{{
    "fill_instructions": [
        {{
            "target_text": "性别：________________",
            "replacement": "性别：男",
            "field_name": "性别",
            "confidence": 0.95
        }}
    ],
    "total_matches": 1
}}
```

【关键要求】
- 必须返回JSON格式，不要返回修改后的文档内容
- target_text必须从原文档中精确复制
- replacement保持原有格式，只替换下划线部分
- 如果没有匹配的数据，该字段就不要包含在指令中"""

        return prompt
    
    def _format_fill_data_for_llm(self, fill_data: Dict[str, Any]) -> str:
        """格式化填充数据供LLM使用"""
        if not fill_data:
            return "（无填充数据）"
        
        formatted_data = []
        for field_name, field_value in fill_data.items():
            if field_value and str(field_value).strip():
                formatted_data.append(f"- {field_name}: {field_value}")
        
        return "\n".join(formatted_data) if formatted_data else "（无有效填充数据）"
    
    def _parse_llm_fill_instructions_with_retry(self, llm_response: str, original_text: str, 
                                               processed_fields: Dict[str, Any], 
                                               form_analysis_fields: List[Dict[str, Any]], 
                                               max_retries: int = 2) -> list:
        """带重试机制的LLM指令解析"""
        # 首先尝试解析原响应
        fill_instructions = self._parse_llm_fill_instructions(llm_response)
        
        if fill_instructions:
            return fill_instructions
        
        # 如果解析失败，进行重试
        for retry_count in range(max_retries):
            self.logger.warning(f"🔄 JSON解析失败，第{retry_count + 1}次重试...")
            
            try:
                # 构建更明确的重试提示
                retry_prompt = self._build_retry_prompt(original_text, processed_fields, form_analysis_fields, retry_count)
                
                # 重新调用LLM
                
                llm_client = LocalLLMClientAdapter()
                retry_response = llm_client.simple_chat(retry_prompt)
                
                if retry_response and retry_response.strip():
                    self.logger.debug(f"🔍 重试{retry_count + 1}响应长度: {len(retry_response)}")
                    
                    # 尝试解析重试响应
                    retry_instructions = self._parse_llm_fill_instructions(retry_response)
                    
                    if retry_instructions:
                        self.logger.info(f"{node_state}-=-填表员===✅ 第{retry_count + 1}次重试成功，获得{len(retry_instructions)}条指令")
                        return retry_instructions
                    else:
                        self.logger.warning(f"❌ 第{retry_count + 1}次重试仍然解析失败")
                else:
                    self.logger.warning(f"❌ 第{retry_count + 1}次重试LLM返回空响应")
                    
            except Exception as e:
                self.logger.error(f"❌ 第{retry_count + 1}次重试异常: {e}")
        
        self.logger.error(f"❌ 经过{max_retries}次重试，仍无法获得有效的JSON指令")
        return []
    
    def _build_retry_prompt(self, original_text: str, processed_fields: Dict[str, Any], 
                           form_analysis_fields: List[Dict[str, Any]], retry_count: int) -> str:
        """构建重试提示，强调JSON格式"""
        base_prompt = self._build_llm_analysis_prompt(original_text, processed_fields, form_analysis_fields)
        
        # 根据重试次数调整提示策略
        if retry_count == 0:
            retry_suffix = """

🚨 CRITICAL: 上一次响应的JSON格式有问题，请务必严格按照以下格式返回：

```json
{
    "fill_instructions": [
        {
            "field_name": "姓名",
            "target_text": "姓名：________________",
            "replacement": "姓名：李汉华",
            "confidence": 0.9
        }
    ]
}
```

请确保：
1. 使用 ```json 代码块包围
2. JSON格式完全正确，没有语法错误
3. 所有字符串都用双引号包围
4. 不要添加任何注释或解释文字
"""
        else:
            retry_suffix = """

🚨 FINAL ATTEMPT: 这是最后一次机会，请返回标准JSON格式：

```json
{"fill_instructions": [{"field_name": "姓名", "target_text": "姓名：________________", "replacement": "姓名：李汉华", "confidence": 0.9}]}
```

严格要求：
- 必须是有效的JSON
- 必须包含 fill_instructions 数组
- 每个指令必须有 field_name, target_text, replacement, confidence 字段
- 不要有任何多余的文字
"""
        
        return base_prompt + retry_suffix
    
    def _parse_llm_fill_instructions(self, llm_response: str) -> list:
        """解析LLM返回的填充指令"""
        try:
            import json
            import re
            
            self.logger.debug("🔍 开始解析LLM响应为填充指令")
            
            # 查找JSON块
            json_pattern = r'```json\s*(.*?)\s*```'
            json_match = re.search(json_pattern, llm_response, re.DOTALL)
            
            if json_match:
                json_text = json_match.group(1)
                self.logger.debug("🔍 找到JSON代码块")
            else:
                # 尝试直接解析整个响应
                json_text = llm_response.strip()
                self.logger.debug("🔍 未找到JSON代码块，尝试解析整个响应")
            
            self.logger.debug(f"🔍 待解析的JSON文本长度: {len(json_text)}")
            self.logger.debug(f"🔍 JSON文本前100字符: {json_text[:100]}")
            
            # 解析JSON
            instructions_data = json.loads(json_text)
            fill_instructions = instructions_data.get('fill_instructions', [])
            
            self.logger.info(f"{node_state}-=-填表员===✅ 成功解析 {len(fill_instructions)} 条填充指令")
            
            # 显示前几条指令的概要
            if fill_instructions:
                for i, instruction in enumerate(fill_instructions[:3]):
                    target = instruction.get('target_text', '')[:50]
                    replacement = instruction.get('replacement', '')[:50]
                    self.logger.debug(f"   指令{i+1}: {target} → {replacement}")
            
            return fill_instructions
            
        except json.JSONDecodeError as e:
            self.logger.error(f"❌ JSON解析失败: {e}")
            self.logger.error(f"💡 LLM可能没有按照JSON格式返回，而是直接返回了修改后的文档")
            self.logger.debug(f"原始响应前500字符: {llm_response[:500]}...")
            
            # 检查是否LLM返回的是修改后的文档而不是JSON
            if '```json' not in llm_response and 'fill_instructions' not in llm_response:
                self.logger.warning("⚠️ LLM似乎返回了修改后的文档而不是JSON指令，这会导致重复填充问题")
            
            return []
        except Exception as e:
            self.logger.error(f"❌ 解析填充指令失败: {str(e)}")
            return []
    
    def _execute_fill_instructions(self, text: str, fill_instructions: list) -> tuple:
        """执行填充指令"""
        try:
            filled_text = text
            filled_count = 0
            
            # 按照置信度排序，优先执行高置信度的指令
            sorted_instructions = sorted(fill_instructions, key=lambda x: x.get('confidence', 0), reverse=True)
            
            for instruction in sorted_instructions:
                target_text = instruction.get('target_text', '')
                replacement = instruction.get('replacement', '')
                field_name = instruction.get('field_name', '')
                confidence = instruction.get('confidence', 0)
                
                if not target_text or not replacement:
                    continue
                
                # 执行替换 - 使用智能模式匹配，忽略下划线数量差异
                replaced = False
                
                # 首先尝试精确匹配
                if target_text in filled_text:
                    filled_text = filled_text.replace(target_text, replacement)
                    filled_count += 1
                    replaced = True
                    self.logger.debug(f"✅ 精确匹配成功: {field_name} (置信度: {confidence:.2f})")
                    self.logger.debug(f"   {target_text} → {replacement}")
                else:
                    # 精确匹配失败，使用智能模式匹配（忽略下划线数量）
                    success, new_text = self._flexible_field_replace(filled_text, field_name, target_text, replacement)
                    if success:
                        filled_text = new_text
                        filled_count += 1
                        replaced = True
                        self.logger.info(f"{node_state}-=-填表员===✅ 智能模式匹配成功: {field_name} (置信度: {confidence:.2f})")
                        self.logger.debug(f"{node_state}-=-填表员===   灵活替换: {field_name} → {replacement}")
                    else:
                        self.logger.warning(f"{node_state}-=-填表员===⚠️ 字段匹配失败: {field_name}")
                        self.logger.debug(f"{node_state}-=-填表员===   目标文本: {target_text}")
                        
                        # 添加更详细的调试信息
                        import re
                        field_pattern = rf'{re.escape(field_name)}\s*[：:]\s*[_\s]*'
                        matches = re.findall(field_pattern, filled_text)
                        if matches:
                            self.logger.debug(f"   文档中找到相似模式: {matches}")
                        else:
                            self.logger.debug(f"   文档中未找到 '{field_name}' 相关模式")
            
            return filled_text, filled_count
            
        except Exception as e:
            self.logger.error(f"执行填充指令失败: {str(e)}")
            return text, 0
    
    def _direct_field_mapping_fill(self, text: str, processed_fields: Dict[str, Any], form_analysis_fields: List[Dict[str, Any]]) -> tuple:
        """基于LightRAG格式化结果进行直接字段匹配填写"""
        filled_count = 0
        modified_text = text
        
        # 临时启用调试级别日志以查看匹配详情
        import logging
        original_level = self.logger.level
        # self.logger.setLevel(logging.)
        
        self.logger.info(f"{node_state}-=-填表员===开始直接字段匹配填写，共有 {len(processed_fields)} 个字段")
        self.logger.debug(f"{node_state}-=-填表员===📄 文档文本长度: {len(text)} 字符")
        
        # 显示文档的前200字符作为参考
        preview = text[:200].replace('\n', '\\n').replace('\r', '\\r')
        self.logger.debug(f"{node_state}-=-填表员===📄 文档开头预览: {preview}...")
        
        # 运行匹配测试（仅在第一次运行时）

        
        # 处理LightRAG返回的格式化字段数据
        for field_name, field_data in processed_fields.items():
            # 兼容两种数据格式
            if isinstance(field_data, dict) and 'value' in field_data:
                # 新格式：{'value': '王明轩', 'label': '姓名', 'type': 'text'}
                field_value = field_data.get('value', '')
                field_label = field_data.get('label', field_name)
            else:
                # 旧格式：直接是值
                field_value = field_data
                field_label = field_name
            
            if not field_value or str(field_value).strip() == '':
                continue
            
            self.logger.info(f"{node_state}-=-填表员===尝试填写字段: {field_name} -> {field_value}")
            
            # 智能字段名称映射，处理LightRAG返回字段名与实际表单字段名的差异
            labels_to_try = []
            
            # 1. 优先使用字段的原始标签（如果有的话）
            if field_label and field_label.strip():
                labels_to_try.append(field_label.strip())
            
            # 2. 使用字段名本身
            if field_name and field_name != field_label:
                labels_to_try.append(field_name)
            
            # 3. 添加智能字段名映射 - 处理常见的名称变体
            field_name_variants = self._get_field_name_variants(field_name, field_label)
            labels_to_try.extend(field_name_variants)
            
            # 尝试填写
            filled = False
            for label in labels_to_try:
                if self._try_direct_replace(modified_text, label, str(field_value)):
                    modified_text = self._perform_direct_replace(modified_text, label, str(field_value))
                    filled_count += 1
                    filled = True
                    self.logger.info(f"{node_state}-=-填表员===✅ 字段填写成功: {label} -> {field_value}")
                    break
            
            if not filled:
                self.logger.warning(f"⚠️ 字段未能填写: {field_name} -> {field_value}, 尝试的标签: {labels_to_try}")
        
        self.logger.info(f"{node_state}-=-填表员===直接字段匹配填写完成，成功填写 {filled_count} 个字段")
        
        # 恢复原始日志级别
        # self.logger.setLevel(original_level)
        
        return modified_text, filled_count
    
    
    def _try_direct_replace(self, text: str, label: str, field_value: str) -> bool:
        """检查是否能直接替换该字段 - 支持多行格式检测"""
        import re
        
        # 1. 多行格式匹配 - 优先检查
        multiline_patterns = [
            rf'{re.escape(label)}\s*[：:]\s*\n\s*_+',               # 标签：\n___
            rf'{re.escape(label)}\s*[：:]\s*\r?\n\s*_+',            # 标签：\r\n___
            rf'{re.escape(label)}\n\s*_+',                         # 标签\n___
            rf'{re.escape(label)}\r?\n\s*_+',                      # 标签\r\n___
            rf'{re.escape(label)}\s*[：:]\s*\n\s*\n\s*_+',          # 标签：\n\n___
        ]
        
        for pattern in multiline_patterns:
            if re.search(pattern, text, re.IGNORECASE | re.MULTILINE):
                self.logger.debug(f"多行匹配成功 - 模式: {pattern}, 标签: {label}")
                return True
        
        # 2. 智能行检查：查找标签附近的独立下划线行
        if self._check_nearby_underscore_lines(text, label):
            self.logger.debug(f"智能行匹配成功 - 标签: {label}")
            return True
        
        # 3. 单行格式匹配
        inline_patterns = [
            # 基本格式：标签：下划线或空格
            rf'{re.escape(label)}\s*[：:]\s*_+',                    # 标签：___
            rf'{re.escape(label)}\s*[：:]\s*\s{{3,}}',              # 标签：空格 (3个或更多)
            rf'{re.escape(label)}\s*[：:]\s*[_\s]{{3,}}',           # 标签：下划线或空格混合
            
            # 无冒号格式
            rf'{re.escape(label)}\s+_+',                           # 标签 ___
            rf'{re.escape(label)}_+',                              # 标签___
            rf'{re.escape(label)}\s+[_\s]{{3,}}',                  # 标签 空格或下划线
            
            # 带括号说明的格式
            rf'{re.escape(label)}\s*[\(（][^）)]*[\)）]\s*[：:]\s*_+',  # 标签(说明)：___
            rf'{re.escape(label)}\s*[\(（][^）)]*[\)）]\s*_+',         # 标签(说明)___
            
            # 更宽松的匹配：标签后面跟着多个特殊字符
            rf'{re.escape(label)}\s*[：:]*\s*[_\-—]+',              # 标签：___--- 或类似
            rf'{re.escape(label)}\s*[：:]*\s*\.{{3,}}',             # 标签：...
            
            # 表格形式：标签|空白区域
            rf'{re.escape(label)}\s*[\|丨]\s*[_\s]{{3,}}',          # 标签|___
        ]
        
        for pattern in inline_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                self.logger.debug(f"单行匹配成功 - 模式: {pattern}, 标签: {label}")
                return True
        
        # 4. 最宽松的匹配
        loose_pattern = rf'{re.escape(label)}.{{0,20}}[_\s\-—]{{5,}}'  # 标签后20字符内有5个或更多填写字符
        if re.search(loose_pattern, text, re.IGNORECASE):
            self.logger.debug(f"宽松匹配成功 - 标签: {label}")
            return True
        
        # 调试：显示匹配失败的上下文
        self._debug_match_failure(text, label)
        return False
        
    def _check_nearby_underscore_lines(self, text: str, label: str) -> bool:
        """检查标签附近是否有独立的下划线行"""
        import re
        
        lines = text.split('\n')
        
        # 找到包含标签的行
        for i, line in enumerate(lines):
            if re.search(re.escape(label), line, re.IGNORECASE):
                # 检查后续几行是否有下划线
                for offset in range(1, 4):  # 检查接下来的3行
                    if i + offset < len(lines):
                        target_line = lines[i + offset]
                        # 如果这一行主要是下划线、空格或横线
                        if re.match(r'^\s*[_\s\-—]{3,}\s*$', target_line):
                            return True
        return False
    
    def _debug_match_failure(self, text: str, label: str):
        """调试匹配失败的情况，显示相关文本上下文"""
        import re
        
        # 查找标签在文本中的位置
        label_matches = list(re.finditer(re.escape(label), text, re.IGNORECASE))
        
        if label_matches:
            self.logger.debug(f"🔍 调试标签匹配失败: '{label}'")
            for i, match in enumerate(label_matches[:3]):  # 只显示前3个匹配
                start = max(0, match.start() - 30)
                end = min(len(text), match.end() + 50)
                context = text[start:end].replace('\n', '\\n').replace('\r', '\\r')
                self.logger.debug(f"  匹配{i+1}: ...{context}...")
        else:
            # 尝试模糊搜索
            similar_pattern = rf'[{re.escape(label[0])}][^{re.escape(label[-1])}]*[{re.escape(label[-1])}]' if len(label) > 1 else re.escape(label)
            similar_matches = list(re.finditer(similar_pattern, text, re.IGNORECASE))
            
            if similar_matches:
                self.logger.debug(f"🔍 未找到完全匹配 '{label}'，但发现相似文本:")
                for i, match in enumerate(similar_matches[:3]):
                    start = max(0, match.start() - 20)
                    end = min(len(text), match.end() + 30)
                    context = text[start:end].replace('\n', '\\n').replace('\r', '\\r')
                    self.logger.debug(f"  相似{i+1}: ...{context}...")
            else:
                self.logger.debug(f"🔍 文档中完全没有找到标签 '{label}' 或相似文本")
    
    def _perform_direct_replace(self, text: str, label: str, field_value: str) -> str:
        """执行直接字段替换 - 支持多行格式和精确定位"""
        import re
        
        # 1. 优先处理多行格式：标签在一行，下划线在下一行
        multiline_patterns = [
            # 标签：\n______ 格式
            (rf'({re.escape(label)}\s*[：:]\s*)\n(\s*)_+', f'\\g<1>\\g<2>{field_value}'),
            (rf'({re.escape(label)}\s*[：:]\s*)\r?\n(\s*)_+', f'\\g<1>\\g<2>{field_value}'),
            
            # 标签\n______ 格式（无冒号）
            (rf'({re.escape(label)})\n(\s*)_+', f'\\g<1>\\n\\g<2>{field_value}'),
            (rf'({re.escape(label)})\r?\n(\s*)_+', f'\\g<1>\\r\\n\\g<2>{field_value}'),
            
            # 标签：\n\n______ 格式（有空行）
            (rf'({re.escape(label)}\s*[：:]\s*)\n\s*\n(\s*)_+', f'\\g<1>\\n\\g<2>{field_value}'),
        ]
        
        # 先尝试多行格式
        for pattern, replacement in multiline_patterns:
            new_text = re.sub(pattern, replacement, text, flags=re.IGNORECASE | re.MULTILINE)
            if new_text != text:
                self.logger.debug(f"多行替换成功 - 模式: {pattern}, 标签: {label} -> {field_value}")
                return new_text
        
        # 2. 单行格式：标签和填写区域在同一行
        inline_patterns = [
            # 基本格式：标签：下划线或空格
            (rf'({re.escape(label)}\s*[：:]\s*)_+', f'\\g<1>{field_value}'),
            (rf'({re.escape(label)}\s*[：:]\s*)\s{{3,}}', f'\\g<1>{field_value}'),
            (rf'({re.escape(label)}\s*[：:]\s*)[_\s]{{3,}}', f'\\g<1>{field_value}'),
            
            # 无冒号格式
            (rf'({re.escape(label)}\s+)_+', f'\\g<1>{field_value}'),
            (rf'({re.escape(label)})_+', f'\\g<1>{field_value}'),
            (rf'({re.escape(label)}\s+)[_\s]{{3,}}', f'\\g<1>{field_value}'),
            
            # 带括号说明的格式
            (rf'({re.escape(label)}\s*[\(（][^）)]*[\)）]\s*[：:]\s*)_+', f'\\g<1>{field_value}'),
            (rf'({re.escape(label)}\s*[\(（][^）)]*[\)）]\s*)_+', f'\\g<1>{field_value}'),
            
            # 特殊字符格式
            (rf'({re.escape(label)}\s*[：:]*\s*)[_\-—]+', f'\\g<1>{field_value}'),
            (rf'({re.escape(label)}\s*[：:]*\s*)\.{{3,}}', f'\\g<1>{field_value}'),
            
            # 表格形式
            (rf'({re.escape(label)}\s*[\|丨]\s*)[_\s]{{3,}}', f'\\g<1>{field_value}'),
        ]
        
        # 尝试单行格式
        for pattern, replacement in inline_patterns:
            new_text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
            if new_text != text:
                self.logger.debug(f"单行替换成功 - 模式: {pattern}, 标签: {label} -> {field_value}")
                return new_text
        
        # 3. 智能查找：查找标签附近的独立下划线行并替换
        new_text = self._smart_replace_nearby_lines(text, label, field_value)
        if new_text != text:
            self.logger.debug(f"智能行替换成功 - 标签: {label} -> {field_value}")
            return new_text
        
        # 4. 最宽松的替换
        loose_pattern = rf'({re.escape(label)}.{{0,20}}?)[_\s\-—]{{5,}}'
        new_text = re.sub(loose_pattern, f'\\g<1>{field_value}', text, flags=re.IGNORECASE)
        if new_text != text:
            self.logger.debug(f"宽松替换成功 - 标签: {label} -> {field_value}")
            return new_text
        
        return text
        
    def _smart_replace_nearby_lines(self, text: str, label: str, field_value: str) -> str:
        """智能替换：查找标签附近的独立下划线行"""
        import re
        
        lines = text.split('\n')
        label_line_indices = []
        
        # 找到包含标签的行
        for i, line in enumerate(lines):
            if re.search(re.escape(label), line, re.IGNORECASE):
                label_line_indices.append(i)
        
        # 对于每个标签行，查找附近的下划线行
        for label_idx in label_line_indices:
            # 检查后续几行是否有下划线
            for offset in range(1, 4):  # 检查接下来的3行
                if label_idx + offset < len(lines):
                    target_line = lines[label_idx + offset]
                    
                    # 如果这一行主要是下划线或空格
                    if re.match(r'^\s*[_\s\-—]{3,}\s*$', target_line):
                        # 替换这一行
                        leading_spaces = re.match(r'^(\s*)', target_line).group(1)
                        lines[label_idx + offset] = f"{leading_spaces}{field_value}"
                        self.logger.debug(f"替换了第{label_idx + offset + 1}行的下划线为: {field_value}")
                        return '\n'.join(lines)
        
        return text
    
 
    
    
    def _fill_word_form_controls(self, doc, filled_fields: Dict[str, Any]) -> int:
        """填写Word表单控件"""
        # 这个功能需要更高级的库支持，目前返回0
        return 0
    
    def _insert_word_content(self, doc, filled_fields: Dict[str, Any]) -> int:
        """在指定位置插入内容"""
        filled_count = 0
        
        # 处理LightRAG字段数据格式
        processed_fields = {}
        for field_name, field_data in filled_fields.items():
            if isinstance(field_data, dict) and 'value' in field_data:
                processed_fields[field_name] = field_data['value']
            else:
                processed_fields[field_name] = str(field_data)
        
        # 如果有填写数据但之前的方法没有成功替换，在文档末尾添加填写信息
        if processed_fields:
            # 添加分隔符
            doc.add_paragraph("\n" + "="*50)
            doc.add_paragraph("📋 LightRAG自动填写信息")
            doc.add_paragraph("="*50)
            
            # 添加填写的字段信息
            for field_name, field_value in processed_fields.items():
                para = doc.add_paragraph()
                para.add_run(f"🔹 {field_name}: ").bold = True
                para.add_run(str(field_value))
                filled_count += 1
            
            # 添加时间戳
            from datetime import datetime
            doc.add_paragraph(f"\n⏰ 填写时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"🤖 填写方法: LightRAG综合推理")
        
        return filled_count
    
    def _flexible_field_replace(self, text: str, field_name: str, target_text: str, replacement: str) -> tuple:
        """
        灵活的字段替换，忽略下划线数量差异
        返回 (是否成功, 替换后的文本)
        """
        try:
            import re
            
            # 从replacement中提取实际的值
            # 例如 "姓名：王明轩" → "王明轩"
            if "：" in replacement:
                field_value = replacement.split("：", 1)[1].strip()
            elif ":" in replacement:
                field_value = replacement.split(":", 1)[1].strip()
            else:
                field_value = replacement.strip()
            
            # 尝试多种灵活的模式匹配，忽略下划线的确切数量
            patterns = [
                # 标准格式：姓名：_____（任意数量下划线）
                rf'({re.escape(field_name)}\s*[：:]\s*)(_{{3,}})',
                # 紧凑格式：姓名_____（任意数量下划线）
                rf'({re.escape(field_name)})(_{{5,}})',
                # 带空格：姓名 _____（任意数量下划线）
                rf'({re.escape(field_name)}\s+)(_{{3,}})',
                # 多行格式：姓名：\n_____
                rf'({re.escape(field_name)}\s*[：:]\s*\n\s*)(_{{3,}})',
                # 更宽松的格式：姓名后面跟任意空白和下划线
                rf'({re.escape(field_name)}\s*[：:]?\s*)(_{{3,}})',
            ]
            
            for i, pattern in enumerate(patterns):
                match = re.search(pattern, text)
                if match:
                    # 找到匹配，进行替换
                    prefix = match.group(1)
                    underscores = match.group(2)
                    
                    # 构建替换文本，保持原有格式
                    if "：" in prefix or ":" in prefix:
                        new_text = prefix + field_value
                    else:
                        # 如果没有冒号，添加冒号
                        new_text = prefix.rstrip() + "：" + field_value
                    
                    # 执行替换
                    result_text = text.replace(match.group(0), new_text)
                    
                    self.logger.debug(f"🔧 灵活模式匹配成功 (模式{i+1}):")
                    self.logger.debug(f"   字段名: {field_name}")
                    self.logger.debug(f"   匹配模式: {pattern}")
                    self.logger.debug(f"   原文本: '{match.group(0)}'")
                    self.logger.debug(f"   新文本: '{new_text}'")
                    
                    return True, result_text
            
            # 如果标准模式都失败，尝试最宽松的匹配
            # 查找字段名后面的任意下划线组合
            loose_pattern = rf'{re.escape(field_name)}[^：:\n\r]*?(_{{3,}})'
            loose_match = re.search(loose_pattern, text)
            if loose_match:
                # 宽松匹配成功，直接替换下划线部分
                old_underscores = loose_match.group(1)
                result_text = text.replace(old_underscores, field_value, 1)  # 只替换第一个匹配
                
                self.logger.debug(f"🔧 宽松模式匹配成功:")
                self.logger.debug(f"   匹配文本: '{loose_match.group(0)}'")
                self.logger.debug(f"   替换下划线: '{old_underscores}' → '{field_value}'")
                
                return True, result_text
            
            # 最后的尝试：如果target_text中包含了正确的字段名，提取出模式再匹配
            if field_name in target_text:
                # 从target_text中提取模式，然后用正则表达式查找
                try:
                    # 将target_text中的下划线替换为正则表达式
                    escaped_target = re.escape(target_text)
                    # 将连续的下划线转换为灵活匹配模式
                    flexible_pattern = re.sub(r'(_+)', r'_+', escaped_target)
                    flexible_pattern = flexible_pattern.replace(r'\_+', r'_+')
                    
                    final_match = re.search(flexible_pattern, text)
                    if final_match:
                        result_text = text.replace(final_match.group(0), replacement)
                        
                        self.logger.debug(f"🔧 目标模式匹配成功:")
                        self.logger.debug(f"   灵活模式: {flexible_pattern}")
                        self.logger.debug(f"   匹配到: '{final_match.group(0)}'")
                        self.logger.debug(f"   替换为: '{replacement}'")
                        
                        return True, result_text
                except Exception as pattern_error:
                    self.logger.debug(f"目标模式匹配失败: {pattern_error}")
            
            return False, text
            
        except Exception as e:
            self.logger.error(f"灵活字段替换异常: {e}")
            return False, text


class ExcelWriter(BaseDocumentWriter):
    """Excel文档写入器"""
    
    def write_document(self, 
                           original_path: Path,
                           filled_data: Dict[str, Any],
                           output_path: Path,
                           form_analysis_result: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """写入Excel文档"""
        try:
            # 检查依赖
            try:
                import openpyxl
                from openpyxl import load_workbook
            except ImportError:
                # 如果没有openpyxl，使用简单复制方法
                return self._create_simple_excel_copy(original_path, filled_data, output_path)
            
            filled_fields = filled_data.get('filled_fields', {})
            
            if not filled_fields:
                return {
                    'success': False,
                    'error': '没有需要填写的数据',
                    'method': 'excel_document'
                }
            
            try:
                # 复制原始文档
                shutil.copy2(original_path, output_path)
                
                # 加载工作簿
                workbook = load_workbook(output_path)
                
                filled_count = 0
                
                # 遍历所有工作表
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    sheet_filled = self._fill_excel_sheet(sheet, filled_fields)
                    filled_count += sheet_filled
                
                # 保存工作簿
                workbook.save(output_path)
                
                return {
                    'success': True,
                    'output_file': str(output_path),
                    'method': 'excel_document',
                    'filled_count': filled_count
                }
                
            except Exception as e:
                self.logger.warning(f"Excel文档处理失败，使用简单复制方法: {e}")
                return self._create_simple_excel_copy(original_path, filled_data, output_path)
            
        except Exception as e:
            self.logger.error(f"Excel文档写入失败: {e}")
            return {
                'success': False,
                'output_file': None,
                'error': str(e),
                'method': 'excel_document'
            }

    def _create_simple_excel_copy(self, 
                                       original_path: Path,
                                       filled_data: Dict[str, Any],
                                       output_path: Path) -> Dict[str, Any]:
        """简单复制Excel文档并添加文本信息"""
        try:
            import shutil
            
            filled_fields = filled_data.get('filled_fields', {})
            
            if not filled_fields:
                return {
                    'success': False,
                    'error': '没有需要填写的数据',
                    'method': 'simple_excel_copy'
                }
            
            # 简单复制原始文档
            output_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(original_path, output_path)
            
            # 创建一个伴随的文本文件包含填写信息
            text_output_path = output_path.with_suffix('.xlsx_filled.txt')
            
            content = f"# Excel文档填写信息\n"
            content += f"# 原始文件: {original_path.name}\n"
            content += f"# 填写时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
            
            for field_name, field_value in filled_fields.items():
                content += f"{field_name}: {field_value}\n"
            
            with open(text_output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return {
                'success': True,
                'output_file': str(output_path),
                'method': 'simple_excel_copy',
                'filled_count': len(filled_fields),
                'text_file': str(text_output_path)
            }
            
        except Exception as e:
            self.logger.error(f"Excel简单复制失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'method': 'simple_excel_copy'
            }
    
    def _fill_excel_sheet(self, sheet, filled_fields: Dict[str, Any]) -> int:
        """填写Excel工作表"""
        filled_count = 0
        
        # 处理LightRAG字段数据格式
        processed_fields = {}
        for field_name, field_data in filled_fields.items():
            if isinstance(field_data, dict) and 'value' in field_data:
                # LightRAG格式：{'value': '值', 'method': '...', ...}
                processed_fields[field_name] = field_data['value']
            else:
                # 简单格式：直接是值
                processed_fields[field_name] = str(field_data)
        
        # 扫描工作表寻找标签和对应的空白单元格
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    cell_text = str(cell.value).lower().strip()
                    
                    # 查找匹配的字段
                    for field_name, field_value in processed_fields.items():
                        # 创建更多可能的标签匹配模式
                        field_labels = [
                            field_name.lower(),
                            f"{field_name.lower()}:",
                            f"{field_name.lower()}：",
                            f"{field_name.lower()})",
                            f"({field_name.lower()})",
                            f"[{field_name.lower()}]",
                            f"{{{field_name.lower()}}}",
                        ]
                        
                        # 检查是否包含占位符模式
                        if any(label in cell_text for label in field_labels):
                            # 首先尝试直接替换占位符
                            original_value = str(cell.value)
                            modified_value = original_value
                            
                            placeholders = [
                                f"{{{field_name}}}",
                                f"[{field_name}]",
                                f"__{field_name}__",
                                f"({field_name})",
                                f"<{field_name}>",
                            ]
                            
                            for placeholder in placeholders:
                                if placeholder in modified_value:
                                    modified_value = modified_value.replace(placeholder, str(field_value))
                                    cell.value = modified_value
                                    filled_count += 1
                                    self.logger.debug(f"Excel直接替换成功: {placeholder} -> {field_value}")
                                    break
                            else:
                                # 如果没有直接替换成功，在右侧或下方单元格填写值
                                target_cell = self._find_excel_target_cell(sheet, cell)
                                if target_cell and not target_cell.value:
                                    target_cell.value = field_value
                                    filled_count += 1
                                    self.logger.debug(f"Excel邻近单元格填写成功: {field_name} -> {field_value}")
                            break
        
        # 如果没有找到合适的位置，在第一个空白区域添加填写信息
        if filled_count == 0 and processed_fields:
            filled_count += self._add_excel_summary(sheet, processed_fields)
        
        return filled_count
    
    def _find_excel_target_cell(self, sheet, label_cell):
        """查找Excel中标签对应的目标单元格"""
        # 尝试右侧单元格
        try:
            right_cell = sheet.cell(
                row=label_cell.row, 
                column=label_cell.column + 1
            )
            if not right_cell.value:
                return right_cell
        except:
            pass
        
        # 尝试下方单元格
        try:
            below_cell = sheet.cell(
                row=label_cell.row + 1, 
                column=label_cell.column
            )
            if not below_cell.value:
                return below_cell
        except:
            pass
        
        return None
    
    def _add_excel_summary(self, sheet, processed_fields: Dict[str, Any]) -> int:
        """在Excel工作表中添加填写信息摘要"""
        filled_count = 0
        
        try:
            # 找到第一个空白区域（通常是右侧或底部）
            max_row = sheet.max_row
            max_col = sheet.max_column
            
            # 在底部添加摘要信息
            start_row = max_row + 2
            
            # 添加标题
            title_cell = sheet.cell(row=start_row, column=1)
            title_cell.value = "📋 LightRAG自动填写信息"
            
            # 设置标题样式（如果可能）
            try:
                from openpyxl.styles import Font, Alignment
                title_cell.font = Font(bold=True, size=12)
                title_cell.alignment = Alignment(horizontal='left')
            except ImportError:
                pass
            
            # 添加分隔线
            separator_cell = sheet.cell(row=start_row + 1, column=1)
            separator_cell.value = "=" * 50
            
            # 添加填写的字段信息
            current_row = start_row + 2
            for field_name, field_value in processed_fields.items():
                # 字段名在第一列
                name_cell = sheet.cell(row=current_row, column=1)
                name_cell.value = f"🔹 {field_name}:"
                
                # 字段值在第二列
                value_cell = sheet.cell(row=current_row, column=2)
                value_cell.value = str(field_value)
                
                current_row += 1
                filled_count += 1
            
            # 添加时间戳
            time_cell = sheet.cell(row=current_row + 1, column=1)
            from datetime import datetime
            time_cell.value = f"⏰ 填写时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            
            method_cell = sheet.cell(row=current_row + 2, column=1)
            method_cell.value = "🤖 填写方法: LightRAG综合推理"
            
            self.logger.info(f"{node_state}-=-填表员===Excel摘要添加完成，添加了{filled_count}个字段信息")
            
        except Exception as e:
            self.logger.warning(f"添加Excel摘要失败: {e}")
        
        return filled_count


class FallbackWriter(BaseDocumentWriter):
    """通用后备写入器 - 为不支持的格式创建文本文件"""
    
    def write_document(self, 
                           original_path: Path,
                           filled_data: Dict[str, Any],
                           output_path: Path,
                           form_analysis_result: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """写入为文本格式"""
        try:
            filled_fields = filled_data.get('filled_fields', {})
            
            if not filled_fields:
                return {
                    'success': False,
                    'error': '没有需要填写的数据',
                    'method': 'fallback_text'
                }
            
            # 修改输出路径为.txt文件
            output_path = output_path.with_suffix('.txt')
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # 创建填写后的文本内容
            content = f"# 填写后的表单数据\n"
            content += f"# 原始文件: {original_path.name}\n"
            content += f"# 填写时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
            
            for field_name, field_value in filled_fields.items():
                content += f"{field_name}: {field_value}\n"
            
            # 尝试读取原始文件内容并添加到末尾
            try:
                if original_path.suffix.lower() == '.txt':
                    with open(original_path, 'r', encoding='utf-8') as f:
                        original_content = f.read()
                    content += f"\n# 原始文件内容:\n{original_content}"
            except Exception as e:
                content += f"\n# 无法读取原始文件内容: {e}"
            
            # 保存文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return {
                'success': True,
                'output_file': str(output_path),
                'method': 'fallback_text',
                'filled_count': len(filled_fields)
            }
            
        except Exception as e:
            self.logger.error(f"后备文本写入失败: {e}")
            return {
                'success': False,
                'output_file': None,
                'error': str(e),
                'method': 'fallback_text'
            }


class TextWriter(BaseDocumentWriter):
    """文本文档写入器"""
    
    def write_document(self, 
                           original_path: Path,
                           filled_data: Dict[str, Any],
                           output_path: Path,
                           form_analysis_result: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """写入文本文档"""
        try:
            filled_fields = filled_data.get('filled_fields', {})
            
            if not filled_fields:
                return {
                    'success': False,
                    'error': '没有需要填写的数据',
                    'method': 'text_document'
                }
            
            # 读取原始文档
            try:
                with open(original_path, 'r', encoding='utf-8') as f:
                    original_content = f.read()
            except UnicodeDecodeError:
                # 尝试其他编码
                with open(original_path, 'r', encoding='gbk') as f:
                    original_content = f.read()
            
            # 创建填写后的内容
            filled_content = original_content
            filled_count = 0
            
            # 简单的占位符替换
            for field_name, field_value in filled_fields.items():
                placeholders = [
                    f"{{{field_name}}}",
                    f"[{field_name}]",
                    f"_{field_name}_",
                    f"<{field_name}>",
                ]
                
                for placeholder in placeholders:
                    if placeholder in filled_content:
                        filled_content = filled_content.replace(placeholder, str(field_value))
                        filled_count += 1
                        break
            
            # 如果没有找到占位符，在文档末尾添加填写信息
            if filled_count == 0:
                filled_content += "\n\n# 填写信息:\n"
                for field_name, field_value in filled_fields.items():
                    filled_content += f"{field_name}: {field_value}\n"
                filled_count = len(filled_fields)
            
            # 保存文档
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(filled_content)
            
            return {
                'success': True,
                'output_file': str(output_path),
                'method': 'text_replacement',
                'filled_count': filled_count
            }
            
        except Exception as e:
            self.logger.error(f"文本文档写入失败: {e}")
            return {
                'success': False,
                'output_file': None,
                'error': str(e),
                'method': 'text_document'
            }